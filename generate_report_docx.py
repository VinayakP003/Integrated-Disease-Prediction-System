from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn

def create_report():
    doc = Document()

    # --- Set Margins (Inside Thesis Spec: Top 3cm, Bottom 3cm, Left 4cm, Right 2cm) ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(3)
        section.left_margin = Cm(4)
        section.right_margin = Cm(2)

    # --- Set Default Style (Times New Roman, 12pt) ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # Set spacing to 1.5
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # --- Title Page ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("INTEGRATED DISEASE PREDICTION SYSTEM:\nA MULTI-MODEL ENSEMBLE APPROACH FOR CARDIOMETABOLIC RISK ASSESSMENT")
    run.bold = True
    run.font.size = Pt(24) # Image said 24 font size for title

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Project Report Submitted in Partial Fulfillment of the Requirements for the Degree of").italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Bachelor of Technology")
    run.bold = True
    run.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("in")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Computer Science & Engineering")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Submitted by:")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Vinayak [Last Name] (Roll No. P003)")
    run.bold = True

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Under the Supervision of:")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[Supervisor Name]\n[Designation]")
    run.bold = True

    doc.add_paragraph("\n\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DIT UNIVERSITY\nMAY, 2026")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_page_break()

    # --- Table of Contents Placeholder ---
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph("[Automatic Table of Contents should be inserted here in MS Word]")
    doc.add_page_break()

    # --- Expanded Chapters ---
    chapters = [
        ("Abstract", "The global healthcare landscape is witnessing a surge in non-communicable diseases (NCDs), particularly those categorized under metabolic syndrome—heart disease, stroke, diabetes, and chronic kidney disease. Traditionally, AI-based diagnostic tools have focused on single-disease prediction, often failing to account for the systemic interplay between these conditions. This project introduces the Integrated Disease Prediction System, a comprehensive clinical decision-support platform.\n\nThe system leverages a Soft Voting Ensemble of machine learning classifiers, including Random Forest, SVM, and Logistic Regression, to predict risks across four disease categories. A novel Cardiometabolic Health Risk Index (CHRI) was developed using meta-learning to provide patients with a consolidated health assessment score. To address the 'Black Box' nature of AI, an Explainable AI (XAI) layer was implemented, providing transparent feedback based on user-provided vitals. Finally, the system bridges the gap between diagnosis and care through a localized Recommendation Engine that discovers real-world Indian healthcare providers. The results demonstrate a high-accuracy, localized, and interpretative tool that empowers preventative healthcare management."),
        
        ("Chapter 1: Introduction", "1.1 Background\nNon-communicable diseases (NCDs) are responsible for over 70% of global deaths. In the Indian context, the rising incidence of sedentary lifestyles and dietary shifts has led to a 'Metabolic Time-bomb.' Conditions like hypertension, obesity, and hyperglycemia often co-exist, significantly increasing the risk of multi-organ failure.\n\n1.2 Problem Statement\nMost existing healthcare applications suffer from three critical flaws: Siloed Diagnostics, Lack of Interpretability, and Geographic Friction. This project addresses these gaps by creating a unified metabolic navigator.\n\n1.3 Objectives\nThe primary objectives are to design a robust multi-disease ensemble prediction engine, formulate the CHRI index, and implement a localized recommendation engine for the Indian market.\n\n1.4 Hardware and Software Requirements\n- Software: Python 3.10+, FastAPI, Scikit-Learn, Pandas, JavaScript (ES6+), HTML5/CSS3.\n- Hardware: Minimum 8GB RAM, i5 Processor (10th Gen or above), 20GB Free Disk Space."),

        ("Chapter 2: Literature Survey", "2.1 Review of Existing Systems\nCurrent research predominantly focuses on isolated organs. While deep learning achieves high accuracy, it lacks the explainability required for clinical trust. Our proposed system bridges this gap by combining high-stability ensembles with a transparent explainability layer.\n\n2.2 Study of Algorithms\n- Random Forest: Used for its robust handling of non-linear medical data and feature importance capabilities.\n- Support Vector Machines (SVM): High-dimensional boundary mapping for clear disease classification.\n- Logistic Regression: Provides well-calibrated probabilities used in our Soft Voting ensemble.\n\n2.3 Technological Trends\nThe rise of Web-scraping and Nominatim APIs allows for a seamless transition from diagnostic output to physical healthcare discovery."),

        ("Chapter 3: System Analysis & Design", "3.1 Software Requirement Specification (SRS)\n- User Input: Age, BMI, Glucose, BP, Hypertension History, City.\n- System Processing: Data cleaning, Model inference, Scraper execution.\n- Output Visualization: 0-100 CHRI score, Risk breakdown bars, Doctor hover-cards.\n\n3.2 System Architecture\nThe system uses a decoupled FastAPI microservices architecture. Each medical condition is handled by an independent ensemble model, while the recommendation service acts as a parallel retrieval engine.\n\n3.3 Algorithm Flow\n1. Receive User Vitals.\n2. Scale data using StandardScaler.\n3. Execute 4 Ensemble Models (Heart, Stroke, Diabetes, CKD).\n4. Compute Meta-score (CHRI).\n5. Extract Feature Importances.\n6. Fetch nearby facilities via Nominatim API.\n7. Scrape specialist data via DuckDuckGo.\n8. Return JSON payload to Glassmorphism Frontend."),

        ("Chapter 4: Implementation Details", "4.1 Backend Services\nThe prediction_service.py manages the loading of joblib models and the execution of the Soft Voting logic. The recommendation_service.py manages the scraper pipeline, including credential cleaning and map-link generation.\n\n4.2 Frontend Visualization\nThe app.js controls the lifecycle of the dashboard. It uses asynchronous fetch calls to populate the bento-grid components. Special attention was paid to 'z-index' management and 'overflow' rules to ensure hover-cards are fully visible over grid boundaries."),

        ("Chapter 5: Testing & Results", "5.1 Test Cases\n| ID | Description | Input | Expected Output | Status |\n|---|---|---|---|---|\n| TC1 | Heart Risk | BP: 150, BMI: 30 | High/Critical Alert | Pass |\n| TC2 | Diabetes Check | Glucose: 190 | Critical Warning | Pass |\n| TC3 | Low Risk Profile | Age: 25, BMI: 22 | Low Score (Green) | Pass |\n| TC4 | Empty Input | None | Validation Error | Pass |\n| TC5 | Location Search | 'Mumbai' | Local Apollo/Max list | Pass |\n\n5.2 Discussion\nThe implementation of SMOTE significantly improved minority class detection in stroke datasets, while the XAI filter improved user engagement by providing 'Honest Explanations' based only on provided data."),

        ("Chapter 6: Conclusion & Future Work", "The Integrated Disease Prediction System successfully bridges the gap between diagnostic math and real-world care. Future iterations will include IoT wearable integration and direct appointment booking APIs."),

        ("Appendix A: Core Backend Implementation", "--- prediction_service.py ---\n[INSERT FULL SOURCE CODE HERE - Adds ~15 pages]\n\n--- recommendation_service.py ---\n[INSERT FULL SOURCE CODE HERE - Adds ~12 pages]"),

        ("Appendix B: Frontend Implementation", "--- styles.css ---\n[INSERT FULL STYLESHEET HERE - Adds ~10 pages]\n\n--- app.js ---\n[INSERT FULL JS LOGIC HERE - Adds ~8 pages]")
    ]

    for title, content in chapters:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(content)
        doc.add_page_break()

    # --- References ---
    doc.add_heading('References', level=1)
    refs = [
        "[1] F. Pedregosa et al., 'Scikit-learn: Machine Learning in Python,' Journal of Machine Learning Research, 2011.",
        "[2] N. V. Chawla et al., 'SMOTE: Synthetic Minority Over-sampling Technique,' Journal of Artificial Intelligence Research, 2002.",
        "[3] World Health Organization, 'Noncommunicable Diseases Country Profiles,' 2024.",
        "[4] OpenStreetMap Contributors, 'Nominatim Search API Documentation,' 2024."
    ]
    for ref in refs:
        doc.add_paragraph(ref)

    # Save
    doc_path = "Integrated_Disease_Prediction_Final_Report.docx"
    doc.save(doc_path)
    print(f"Report generated successfully: {doc_path}")

if __name__ == "__main__":
    create_report()
